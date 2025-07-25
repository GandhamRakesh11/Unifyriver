
#I have divided this file into extract_text, extract_table and describe_image
import os
os.environ.setdefault("HOME", "/app")  # 🔄 Add this
os.environ.setdefault("PADDLEOCR_HOME", "/app/.paddleocr")
import fitz  # PyMuPDF
import numpy as np
from PIL import Image
import pdfplumber
import pandas as pd
from docx.api import Document
import cv2
import tempfile
import re
import logging
import requests
import base64
from paddleocr import PaddleOCR
from requests.auth import HTTPBasicAuth
from collections import Counter
# ✅ Gemini
from google.generativeai import upload_file, GenerativeModel, configure

from dotenv import load_dotenv

load_dotenv()

# ✅ Mistral (your preferred usage)
from mistralai import Mistral

# Optional MistralOCR
try:
    from doctr.models import ocr_predictor
    from doctr.io import DocumentFile
    use_mistral = True
except ImportError:
    use_mistral = False



ocr = PaddleOCR(use_angle_cls=True, lang='en')

mistral_ocr = ocr_predictor(pretrained=True) if use_mistral else None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def compute_caption_metrics(text):
    length_score = len(text.split())
    lower = text.lower()

    density_terms = ["axis", "legend", "trend", "outlier", "distribution", "variance", "correlat", "increase", "decrease"]
    density_score = sum(lower.count(term) for term in density_terms)

    words = text.split()
    word_counts = Counter(words)
    redundancy_score = sum(c for w, c in word_counts.items() if c > 1) / max(len(words), 1)

    noun_like = [w for w in words if w[0].isupper()]
    specificity_score = len(set(noun_like)) / max(len(words), 1)

    insight_keywords = ["trend", "increase", "decrease", "anomaly", "correlat", "cluster", "outlier", "distribution"]
    insight_score = sum(lower.count(k) for k in insight_keywords) / max(len(words), 1)

    hallucination_risk = length_score > 100 and "no graph" in lower

    return {
        "length_score": length_score,
        "density_score": density_score,
        "redundancy_score": redundancy_score,
        "specificity_score": specificity_score,
        "insight_score": insight_score,
        "hallucination_risk": hallucination_risk
    }

def auto_rotate_image(pil_img):
    img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2GRAY)
    coords = np.column_stack(np.where(img_cv > 0))
    angle = cv2.minAreaRect(coords)[-1]
    angle = -(90 + angle) if angle < -45 else -angle
    (h, w) = img_cv.shape[:2]
    M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
    rotated = cv2.warpAffine(img_cv, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    return Image.fromarray(cv2.cvtColor(rotated, cv2.COLOR_GRAY2RGB))

def extract_images_with_fitz(pdf_path):
    doc = fitz.open(pdf_path)
    images = [
        Image.frombytes("RGB", [p.get_pixmap(matrix=fitz.Matrix(2, 2)).width,
                                p.get_pixmap(matrix=fitz.Matrix(2, 2)).height],
                        p.get_pixmap(matrix=fitz.Matrix(2, 2)).samples)
        for p in doc
    ]
    doc.close()
    return images

def extract_from_docx(path, start_page=None, end_page=None, extract_captions=False):
    doc = Document(path)

    # Extract paragraphs
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    page_texts = []
    page_size = 500
    for i in range(0, len(paras), page_size):
        page_texts.append("\n".join(paras[i:i + page_size]))

    selected_pages = page_texts
    if start_page and end_page:
        selected_pages = page_texts[start_page - 1:end_page]

    # Extract tables
    tables = []
    for t in doc.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in t.rows]
        tables.append("\n".join(rows))

    # Extract and caption images if requested
    captions = ""
    if extract_captions:
        try:
            images = []
            for rel in doc.part._rels:
                rel = doc.part._rels[rel]
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                        tmp.write(image_data)
                        tmp.flush()
                        img = Image.open(tmp.name)
                        images.append(img)

            captions = "\n\n".join(
                f"Image {i + 1}:\n{describe_image(img, i + 1)}"
                for i, img in enumerate(images)
                if is_visually_dense(img) or image_entropy(img)
            )
        except Exception as e:
            logging.warning(f"Captioning error for DOCX file: {e}")
            captions = "[Captioning failed for DOCX file]"

    return clean_text("\n\n".join(selected_pages)), "\n\n".join(tables), captions


def extract_from_excel(path):
    xl = pd.ExcelFile(path)
    tables = [f"Sheet: {s}\n{xl.parse(s).to_string(index=False)}" for s in xl.sheet_names]
    return "", "\n\n".join(tables)

def extract_from_csv(path):
    return "", pd.read_csv(path).to_string(index=False)

def is_visually_dense(image, threshold=0.02):
    edges = cv2.Canny(cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY), 100, 200)
    return (np.count_nonzero(edges) / edges.size) > threshold

def image_entropy(image, threshold=4.0):
    gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
    hist = cv2.calcHist([gray], [0], None, [256], [0, 256]).ravel()
    hist /= hist.sum()
    entropy = -np.sum(hist[hist > 0] * np.log2(hist[hist > 0]))
    return entropy > threshold

def upload_to_imagekit(file_path):
    try:
        with open(file_path, 'rb') as f:
            file_data = base64.b64encode(f.read()).decode()

        url = 'https://upload.imagekit.io/api/v1/files/upload'
        payload = {
            'file': file_data,
            'fileName': 'mistral_image.png',
            'useUniqueFileName': 'true'
        }

        PRIVATE_API_KEY = 'private_mDVE6rmy2NS2regbjNNZI9QAJO4='
        response = requests.post(url, data=payload, auth=HTTPBasicAuth(PRIVATE_API_KEY, ''))

        if response.status_code == 200:
            return response.json().get('url')
        else:
            logger.error(f"ImageKit upload failed: {response.text}")
            return None
    except Exception as e:
        logger.error(f"ImageKit upload exception: {e}")
        return None

def describe_image(img, page_number):
    gemini_text = "[Gemini not executed]"
    gemini_metrics = {}
    mistral_text = "[Mistral response unavailable]"
    mistral_metrics = {}

    try:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            img.save(tmp.name)
            local_path = tmp.name

        gemini_prompt = (
            "From this image, provide a deep and detailed analysis of all graphs, charts, and visual data present. "
            "For each chart or graph, please provide the following insights:\n\n"
            "1. Type of chart (e.g., bar, line, pie, scatter, heatmap, etc.)\n"
            "2. Title or caption (if present)\n"
            "3. Axes labels and units (X and Y axes)\n"
            "4. Legend or categories involved\n"
            "5. Trends, correlations, or outliers you observe\n"
            "6. Comparison of multiple data series (if applicable)\n"
            "7. Statistical observations such as peaks, averages, distributions, or anomalies\n"
            "8. What the chart reveals about the subject matter or overall context\n"
            "9. Any possible misinterpretation, bias, or limitations in the visual representation\n\n"
            "Structure your analysis clearly by separating each chart, and use bullet points or tables where appropriate "
            "to improve readability. Summarize key insights across all visuals at the end if multiple charts are present."
        )

        # 🧠 Gemini Analysis
        try:
            configure(api_key=os.getenv("GOOGLE_API_KEY"))
            gemini_model = GenerativeModel("gemini-1.5-flash")
            uploaded = upload_file(local_path, mime_type="image/png")
            gemini_response = gemini_model.generate_content([uploaded, gemini_prompt])
            gemini_text = clean_text(gemini_response.text)
            gemini_metrics = compute_caption_metrics(gemini_text)
        except Exception as ge:
            logger.error(f"[Page {page_number}] Gemini error: {ge}")
            gemini_text = f"[Gemini Error: {ge}]"

        # 🎯 Mistral Analysis
        mistral_key = os.getenv("MISTRAL_API_KEY")
        if mistral_key:
            try:
                mistral_client = Mistral(api_key=mistral_key)
                image_url = upload_to_imagekit(local_path)
                if not image_url:
                    mistral_text = "[Image upload failed — invalid URL]"
                else:
                    mistral_response = mistral_client.chat.complete(
                        model="pixtral-12b-latest",
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": gemini_prompt},
                                    {"type": "image_url", "image_url": image_url}
                                ]
                            }
                        ]
                    )
                    mistral_text = mistral_response.choices[0].message.content.strip()
                    mistral_metrics = compute_caption_metrics(mistral_text)
            except Exception as me:
                logger.error(f"[Page {page_number}] Mistral error: {me}")
                mistral_text = f"[Mistral Error: {me}]"
        else:
            mistral_text = "[Mistral API key not set]"

        os.remove(local_path)

        result = f"""[Gemini]
{gemini_text}
[Gemini Metrics]
{gemini_metrics}
[Mistral]
{mistral_text}
[Mistral Metrics]
{mistral_metrics}"""
        return result

    except Exception as e:
        logger.error(f"[Page {page_number}] Caption error: {e}")
        return f"[Caption Error: {e}]"


# (rest of the code remains unchanged)


def extract_from_pdf(path, start_page=None, end_page=None, extract_captions=False):
    result = {"text": [], "tables": [], "captions": []}
    images = extract_images_with_fitz(path)
    doc = fitz.open(path)
    total_pages = len(doc)
    start = max(start_page or 1, 1)
    end = min(end_page or total_pages, total_pages)

    for i, page in enumerate(doc):
        page_num = i + 1
        if page_num < start or page_num > end:
            continue

        text = page.get_text()
        if text.strip():
            result["text"].append(f"Page {page_num} (Extracted):\n{clean_text(text)}")
        else:
            if i < len(images):
                img = auto_rotate_image(images[i])
                img_np = np.array(img)
                try:
                    ocr_result = ocr.ocr(img_np, cls=True)
                    ocr_text = "\n".join([line[1][0] for line in ocr_result[0]]) if ocr_result else ""
                    if not ocr_text and use_mistral:
                        doc_img = DocumentFile.from_images(img)
                        ocr_text = mistral_ocr(doc_img).render()
                except Exception as e:
                    logger.warning(f"OCR error page {page_num}: {e}")
                    ocr_text = "[OCR Error]"
                result["text"].append(f"Page {page_num} (OCR):\n{clean_text(ocr_text) or '[No OCR Text]'}")
            else:
                result["text"].append(f"Page {page_num}:\n[No text or image]")

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_num = i + 1
            if page_num < start or page_num > end:
                continue
            for table in page.extract_tables():
                rows = [" | ".join(cell or "" for cell in row) for row in table if row]
                result["tables"].append(f"Page {page_num} Table:\n" + "\n".join(rows))

    if extract_captions:
        for i, page in enumerate(doc):
            page_num = i + 1
            if page_num < start or page_num > end:
                continue
            if i >= len(images):
                continue
            img = images[i]
            if (page.get_images(full=True) or len(page.get_text().strip()) < 30
                    or is_visually_dense(img) or image_entropy(img)):
                result["captions"].append(f"Page {page_num}:\n" + describe_image(img, page_num))

    doc.close()
    return "\n\n".join(result["text"]), "\n\n".join(result["tables"]), "\n\n".join(result["captions"])

def extract_file(file, start_page=None, end_page=None, extract_captions=False, filename=None):
    ext = os.path.splitext(filename or "")[-1].lower()

    if ext == ".pdf":
        return extract_from_pdf(
            file.name,
            start_page=start_page,
            end_page=end_page,
            extract_captions=extract_captions
        )

    elif ext == ".docx":
        return extract_from_docx(
            file.name,
            start_page=start_page,
            end_page=end_page,
            extract_captions=extract_captions
        )

    elif ext == ".csv":
        text, tables = extract_from_csv(file.name)
        return text, tables, ""

    elif ext in [".xls", ".xlsx"]:
        text, tables = extract_from_excel(file.name)
        return text, tables, ""

    else:
        return "Unsupported file type", "", ""